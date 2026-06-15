"""
Resume text extraction from PDF and DOCX files.
Uses PyMuPDF for PDFs and python-docx for Word documents.
"""

import io
from pathlib import Path


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file given its raw bytes."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file given its raw bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type from filename and extract text.
    Supports PDF and DOCX formats.
    """
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name.endswith(".doc"):
        raise ValueError("Old .doc format is not supported. Please convert to .docx or .pdf.")
    elif name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format: {Path(filename).suffix}")
